from __future__ import annotations

import re
from pathlib import Path

import pythoncom
import win32com.client
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
SOURCE_PPTX = ROOT / "랩미팅발표자료_withcodex.pptx"
OUTPUT_DOCX = ROOT / "script.docx"


def _clean_text(text: str) -> str:
    text = str(text or "").replace("\r", "\n")
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    return "\n".join(lines).strip()


def _is_noise_text(text: str) -> bool:
    if not text:
        return True
    if re.fullmatch(r"\d+", text):
        return True
    lowered = text.strip().lower()
    if lowered in {"yongsang an", "psid lab", "keywords"}:
        return True
    if "20260329" in text or "march 29, 2026" in lowered or "weekly lab meeting" in lowered:
        return True
    return False


def _candidate_titles(slide) -> list[str]:
    items: list[str] = []
    for shape in slide.Shapes:
        try:
            if shape.HasTextFrame and shape.TextFrame.HasText:
                txt = _clean_text(shape.TextFrame.TextRange.Text)
                if txt and not _is_noise_text(txt):
                    items.append(txt)
        except Exception:
            continue
    deduped: list[str] = []
    seen: set[str] = set()
    for item in items:
        key = item.strip()
        if key not in seen:
            seen.add(key)
            deduped.append(key)
    return deduped


def _extract_note_text(slide) -> str:
    parts: list[str] = []
    for shape in slide.NotesPage.Shapes:
        try:
            if shape.HasTextFrame and shape.TextFrame.HasText:
                txt = _clean_text(shape.TextFrame.TextRange.Text)
                if not txt or re.fullmatch(r"\d+", txt):
                    continue
                parts.append(txt)
        except Exception:
            continue
    out: list[str] = []
    seen: set[str] = set()
    for part in parts:
        if part not in seen:
            seen.add(part)
            out.append(part)
    return "\n\n".join(out).strip()


def _set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    style.font.size = Pt(11)


def _configure_sections(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)


def _write_doc(entries: list[dict[str, str]]) -> None:
    doc = Document()
    _set_default_style(doc)
    _configure_sections(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("랩미팅 발표자료 발표 대본\nLab Meeting Presentation Script")
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Source: 랩미팅발표자료_withcodex.pptx\nGenerated automatically from PowerPoint speaker notes")
    r2.font.name = "Calibri"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r2.font.size = Pt(10)

    doc.add_paragraph("")

    for idx, entry in enumerate(entries, start=1):
        heading = doc.add_paragraph()
        heading.style = "Heading 1"
        hr = heading.add_run(f"Slide {idx}. {entry['heading']}")
        hr.font.name = "Calibri"
        hr._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        hr.font.size = Pt(14)
        hr.bold = True

        if entry.get("subtitle"):
            p = doc.add_paragraph()
            p.style = "Heading 2"
            rr = p.add_run(entry["subtitle"])
            rr.font.name = "Calibri"
            rr._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
            rr.font.size = Pt(11)
            rr.bold = True

        body = entry.get("notes") or "(No speaker notes found.)"
        for block in body.split("\n\n"):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.line_spacing = 1.25
            para.paragraph_format.space_after = Pt(6)
            rb = para.add_run(block.strip())
            rb.font.name = "Calibri"
            rb._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
            rb.font.size = Pt(11)

        if idx != len(entries):
            doc.add_page_break()

    doc.save(str(OUTPUT_DOCX))


def main() -> None:
    pythoncom.CoInitialize()
    app = None
    pres = None
    try:
        app = win32com.client.DispatchEx("PowerPoint.Application")
        pres = app.Presentations.Open(str(SOURCE_PPTX), False, True, False)
        entries: list[dict[str, str]] = []
        for i in range(1, pres.Slides.Count + 1):
            slide = pres.Slides.Item(i)
            texts = _candidate_titles(slide)
            heading = texts[1] if len(texts) >= 2 else (texts[0] if texts else f"Slide {i}")
            subtitle = texts[0] if len(texts) >= 2 else ""
            entries.append(
                {
                    "heading": heading,
                    "subtitle": subtitle,
                    "notes": _extract_note_text(slide),
                }
            )
        _write_doc(entries)
        print(f"Saved: {OUTPUT_DOCX}")
    finally:
        if pres is not None:
            pres.Close()
        if app is not None:
            app.Quit()
        pythoncom.CoUninitialize()


if __name__ == "__main__":
    main()
