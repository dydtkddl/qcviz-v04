from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
SOURCE_MD = ROOT / "랩미팅발표_전면개편_완성대본_2026-03-30.md"
OUTPUT_DOCX = ROOT / "랩미팅발표_전면개편_완성대본_2026-03-30.docx"


def build_docx() -> None:
    text = SOURCE_MD.read_text(encoding="utf-8", errors="replace")
    lines = text.splitlines()

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Malgun Gothic"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.style = "Title"
    title_run = title.add_run("랩미팅 발표 전면 개편 완성 대본")
    title_run.font.name = "Malgun Gothic"

    subtitle = doc.add_paragraph()
    subtitle.style = "Subtitle"
    subtitle_run = subtitle.add_run("Generated from markdown source")
    subtitle_run.font.name = "Malgun Gothic"

    current_mode = None

    for raw_line in lines[2:]:
        line = raw_line.rstrip()

        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.startswith("## "):
            p = doc.add_paragraph()
            p.style = "Heading 1"
            run = p.add_run(line[3:].strip())
            run.font.name = "Malgun Gothic"
            current_mode = None
            continue

        if line.startswith("### "):
            heading = line[4:].strip()
            p = doc.add_paragraph()
            p.style = "Heading 2"
            run = p.add_run(heading)
            run.font.name = "Malgun Gothic"
            current_mode = heading
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line[2:].strip())
            run.font.name = "Malgun Gothic"
            continue

        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Malgun Gothic"
        if current_mode == "발표 스크립트":
            run.font.size = Pt(11.5)

    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_docx()
