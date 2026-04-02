from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
SOURCE_MD = ROOT / "QCViz_MCP_Master_Specification_2026-03-30.md"
OUTPUT_DOCX = ROOT / "QCViz_MCP_Master_Specification_2026-03-30.docx"


def set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    style.font.size = Pt(11)


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    add_page_number(section)


def apply_heading_style(paragraph, level: int) -> None:
    if level == 1:
        paragraph.style = "Heading 1"
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
            run.font.size = Pt(18)
            run.bold = True
    elif level == 2:
        paragraph.style = "Heading 2"
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
            run.font.size = Pt(14)
            run.bold = True


def add_markdown(document: Document, text: str) -> None:
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line:
            document.add_paragraph("")
            continue
        if line.startswith("# "):
            p = document.add_paragraph(line[2:].strip())
            apply_heading_style(p, 1)
            continue
        if line.startswith("## "):
            p = document.add_paragraph(line[3:].strip())
            apply_heading_style(p, 2)
            continue
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt = p.paragraph_format
        fmt.line_spacing = 1.3
        fmt.space_after = Pt(6)
        run = p.add_run(line)
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.size = Pt(11)


def build_document() -> Document:
    document = Document()
    set_default_font(document)
    configure_page(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("QCViz-MCP v3\nMaster Specification and Contribution Narrative")
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    title_run.font.size = Pt(20)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Korean and English Full-Length Project Description\nMarch 30, 2026")
    subtitle_run.font.name = "Calibri"
    subtitle_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    subtitle_run.font.size = Pt(11)

    document.add_page_break()
    add_markdown(document, SOURCE_MD.read_text(encoding="utf-8"))

    section = document.add_section(WD_SECTION_START.CONTINUOUS)
    add_page_number(section)
    return document


def main() -> None:
    document = build_document()
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)
    print(f"Saved: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
